"""
Live enterprise connectors — replaces stub_connectors.py.

Connectors:
  ConfluenceConnector  — Atlassian Cloud REST API (atlassian-python-api)
  JiraConnector        — Atlassian Cloud REST API (atlassian-python-api)
  OutlookConnector     — Microsoft Graph API (msal + httpx)
  SharePointConnector  — Microsoft Graph API (msal + httpx)

All implement BaseConnector.fetch() → list[RawDocument].
The existing IngestionPipeline needs zero changes.

Credentials are read from src/config.settings (sourced from .env).
If a required credential is missing, fetch() raises ValueError with a
specific message listing the missing env vars.
"""

import io
import logging
import uuid
from typing import Optional

import httpx

from src.ingestion.connectors.base import BaseConnector, RawDocument
from src.config import settings
from src.storage.database import get_connector_sync_state

logger = logging.getLogger(__name__)

# Retry decorator — 3 attempts, exponential backoff 2→ 4→ 8s
# Retries on network errors and HTTP 5xx; stops immediately on ValueError (bad credentials)
try:
    from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
    _RETRY = retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=8),
        retry=retry_if_exception_type((httpx.HTTPStatusError, httpx.RequestError, RuntimeError)),
        reraise=True,
    )
except ImportError:
    # tenacity not installed — no-op decorator
    def _RETRY(fn):  # type: ignore
        return fn


# ── Helpers ───────────────────────────────────────────────────────────────────

def _html_to_text(html: str) -> str:
    """Convert HTML string to plain text. Falls back gracefully if html2text
    is not installed."""
    try:
        import html2text
        h = html2text.HTML2Text()
        h.ignore_links = True
        h.ignore_images = True
        h.body_width = 0
        return h.handle(html).strip()
    except ImportError:
        # Minimal fallback — strip tags
        import re
        text = re.sub(r"<[^>]+>", " ", html)
        return " ".join(text.split())


# ── Confluence Connector ──────────────────────────────────────────────────────

class ConfluenceConnector(BaseConnector):
    """
    Fetches Confluence pages from a given space via Atlassian Cloud REST API.

    Required .env vars:
        CONFLUENCE_URL   — e.g. https://yourorg.atlassian.net
        CONFLUENCE_USER  — your Atlassian account email
        CONFLUENCE_API_TOKEN — API token from id.atlassian.com

    Optional constructor args:
        space_key        — Confluence space key (e.g. "OPS")
        max_pages        — max pages to fetch (default 500)
        allowed_roles    — RBAC roles (default: all roles)
        sensitivity_level — sensitivity tag (default: "internal")
    """

    def __init__(
        self,
        space_key: str,
        max_pages: int = 500,
        allowed_roles: Optional[list[str]] = None,
        sensitivity_level: str = "internal",
        **kwargs,
    ):
        self.space_key = space_key
        self.max_pages = max_pages
        self.allowed_roles = allowed_roles or ["analyst", "operations", "compliance", "admin"]
        self.sensitivity_level = sensitivity_level

    def _check_credentials(self):
        missing = [
            v for v, val in [
                ("CONFLUENCE_URL", settings.confluence_url),
                ("CONFLUENCE_USER", settings.confluence_user),
                ("CONFLUENCE_API_TOKEN", settings.confluence_api_token),
            ] if not val
        ]
        if missing:
            raise ValueError(
                f"ConfluenceConnector: missing env vars: {', '.join(missing)}. "
                "Add them to your .env file."
            )

    @_RETRY
    def fetch(self) -> list[RawDocument]:
        self._check_credentials()
        try:
            from atlassian import Confluence
        except ImportError:
            raise ImportError(
                "atlassian-python-api is not installed. "
                "Run: pip install atlassian-python-api"
            )

        confluence = Confluence(
            url=settings.confluence_url,
            username=settings.confluence_user,
            password=settings.confluence_api_token,
            cloud=True,
        )

        # ── Incremental Sync Logic ──
        last_sync = get_connector_sync_state(f"confluence:{self.space_key}")
        cql = f'space = "{self.space_key}"'
        if last_sync and last_sync.get("last_synced"):
            # Confluence CQL uses YYYY-MM-DD or relative. We convert ISO to YYYY-MM-DD
            date_str = last_sync["last_synced"].split("T")[0]
            cql += f' AND lastModified > "{date_str}"'
            logger.info("ConfluenceConnector: performing incremental sync since %s", date_str)

        documents: list[RawDocument] = []
        start = 0
        limit = 50  # page size

        logger.info("ConfluenceConnector: fetching from space '%s'", self.space_key)

        while len(documents) < self.max_pages:
            # Use CQL for incremental search if supported, else fallback to space pages
            results = confluence.cql(
                cql,
                start=start,
                limit=limit,
                expand="body.storage,version,space",
            ).get("results", [])

            if not results:
                break

            for page in results:
                try:
                    page_id = page.get("id", "")
                    title = page.get("title", "Untitled")
                    html_body = page.get("body", {}).get("storage", {}).get("value", "")
                    plain_text = _html_to_text(html_body)
                    version = page.get("version", {}).get("number", 1)
                    url = f"{settings.confluence_url}/wiki/spaces/{self.space_key}/pages/{page_id}"

                    if not plain_text.strip():
                        continue

                    doc = RawDocument(
                        source_id=f"confluence-{page_id}",
                        source_name=f"Confluence: {title}",
                        source_path=url,
                        source_type="confluence",
                        content=plain_text,
                        metadata={
                            "page_id": page_id,
                            "space_key": self.space_key,
                            "title": title,
                            "version": version,
                            "url": url,
                        },
                        pages=[{"page_num": 1, "text": plain_text}],
                        allowed_roles=self.allowed_roles,
                        sensitivity_level=self.sensitivity_level,
                    )
                    documents.append(doc)
                except Exception as exc:
                    logger.warning("Confluence page error: %s", exc)

            if len(results) < limit:
                break
            start += limit

        logger.info("ConfluenceConnector: fetched %d pages", len(documents))
        return documents


# ── Jira Connector ────────────────────────────────────────────────────────────

class JiraConnector(BaseConnector):
    """
    Fetches Jira issues (tickets + comments) from a given project.

    Required .env vars:
        JIRA_URL         — e.g. https://yourorg.atlassian.net
        JIRA_USER        — your Atlassian account email
        JIRA_API_TOKEN   — API token from id.atlassian.com

    Optional constructor args:
        project_key      — Jira project key (e.g. "BANK")
        jql              — override default JQL query
        max_issues       — max issues to fetch (default 500)
        include_comments — whether to append comments (default True)
        allowed_roles / sensitivity_level
    """

    def __init__(
        self,
        project_key: str,
        jql: Optional[str] = None,
        max_issues: int = 500,
        include_comments: bool = True,
        allowed_roles: Optional[list[str]] = None,
        sensitivity_level: str = "internal",
        **kwargs,
    ):
        self.project_key = project_key
        self.jql = jql or f'project = "{project_key}" ORDER BY updated DESC'
        self.max_issues = max_issues
        self.include_comments = include_comments
        self.allowed_roles = allowed_roles or ["analyst", "operations", "compliance", "admin"]
        self.sensitivity_level = sensitivity_level

    def _check_credentials(self):
        missing = [
            v for v, val in [
                ("JIRA_URL", settings.jira_url),
                ("JIRA_USER", settings.jira_user),
                ("JIRA_API_TOKEN", settings.jira_api_token),
            ] if not val
        ]
        if missing:
            raise ValueError(
                f"JiraConnector: missing env vars: {', '.join(missing)}. "
                "Add them to your .env file."
            )

    @_RETRY
    def fetch(self) -> list[RawDocument]:
        self._check_credentials()
        try:
            from atlassian import Jira
        except ImportError:
            raise ImportError(
                "atlassian-python-api is not installed. "
                "Run: pip install atlassian-python-api"
            )

        jira = Jira(
            url=settings.jira_url,
            username=settings.jira_user,
            password=settings.jira_api_token,
            cloud=True,
        )

        # ── Incremental Sync Logic ──
        last_sync = get_connector_sync_state(f"jira:{self.project_key}")
        search_jql = self.jql
        if last_sync and last_sync.get("last_synced"):
            # ISO timestamp: 2024-03-25T12:34:56 -> 2024-03-25 12:34
            date_str = last_sync["last_synced"].replace("T", " ")[:16]
            search_jql += f' AND updated >= "{date_str}"'
            logger.info("JiraConnector: performing incremental sync since %s", date_str)

        logger.info("JiraConnector: JQL = %s", search_jql)

        documents: list[RawDocument] = []
        start = 0
        page_size = 50

        while len(documents) < self.max_issues:
            response = jira.jql(
                search_jql,
                start=start,
                limit=page_size,
                fields=["summary", "description", "status", "priority",
                        "assignee", "reporter", "issuetype", "comment",
                        "resolution", "created", "updated"],
            )
            issues = response.get("issues", [])
            if not issues:
                break

            for issue in issues:
                try:
                    key = issue.get("key", "")
                    fields = issue.get("fields", {})

                    summary = fields.get("summary", "")
                    description = _html_to_text(fields.get("description") or "")
                    status = (fields.get("status") or {}).get("name", "")
                    priority = (fields.get("priority") or {}).get("name", "")
                    assignee = ((fields.get("assignee") or {}).get("displayName") or "Unassigned")
                    issue_type = (fields.get("issuetype") or {}).get("name", "")
                    resolution = ((fields.get("resolution") or {}).get("name") or "Unresolved")

                    # Build full text: summary + description + comments
                    parts = [
                        f"Issue: {key}",
                        f"Type: {issue_type}",
                        f"Status: {status}  |  Priority: {priority}",
                        f"Assignee: {assignee}",
                        f"Resolution: {resolution}",
                        f"\nSummary: {summary}",
                    ]
                    if description:
                        parts.append(f"\nDescription:\n{description}")

                    if self.include_comments:
                        comments = (fields.get("comment") or {}).get("comments", [])
                        for c in comments:
                            author = (c.get("author") or {}).get("displayName", "?")
                            body = _html_to_text(c.get("body") or "")
                            parts.append(f"\n[Comment by {author}]:\n{body}")

                    full_text = "\n".join(parts)
                    url = f"{settings.jira_url}/browse/{key}"

                    doc = RawDocument(
                        source_id=f"jira-{key}",
                        source_name=f"Jira {key}: {summary[:60]}",
                        source_path=url,
                        source_type="jira",
                        content=full_text,
                        metadata={
                            "issue_key": key,
                            "project_key": self.project_key,
                            "status": status,
                            "priority": priority,
                            "issue_type": issue_type,
                            "url": url,
                        },
                        pages=[{"page_num": 1, "text": full_text}],
                        allowed_roles=self.allowed_roles,
                        sensitivity_level=self.sensitivity_level,
                    )
                    documents.append(doc)
                except Exception as exc:
                    logger.warning("Jira issue error: %s", exc)

            if len(issues) < page_size:
                break
            start += page_size

        logger.info("JiraConnector: fetched %d issues", len(documents))
        return documents


# ── Microsoft Graph helper ────────────────────────────────────────────────────

def _get_ms_graph_token() -> str:
    """Acquire an app-only Microsoft Graph token via client_credentials flow."""
    try:
        import msal
    except ImportError:
        raise ImportError("msal is not installed. Run: pip install msal")

    missing = [
        v for v, val in [
            ("MS_TENANT_ID", settings.ms_tenant_id),
            ("MS_CLIENT_ID", settings.ms_client_id),
            ("MS_CLIENT_SECRET", settings.ms_client_secret),
        ] if not val
    ]
    if missing:
        raise ValueError(
            f"Microsoft Graph: missing env vars: {', '.join(missing)}. "
            "Add them to your .env file."
        )

    authority = f"https://login.microsoftonline.com/{settings.ms_tenant_id}"
    app = msal.ConfidentialClientApplication(
        settings.ms_client_id,
        authority=authority,
        client_credential=settings.ms_client_secret,
    )
    result = app.acquire_token_for_client(
        scopes=["https://graph.microsoft.com/.default"]
    )
    if "access_token" not in result:
        raise RuntimeError(
            f"Microsoft Graph auth failed: {result.get('error_description', result)}"
        )
    return result["access_token"]


# ── Outlook Connector ─────────────────────────────────────────────────────────

class OutlookConnector(BaseConnector):
    """
    Fetches emails from an Outlook/Exchange mailbox via Microsoft Graph API.

    Required .env vars:
        MS_TENANT_ID, MS_CLIENT_ID, MS_CLIENT_SECRET
        MS_MAILBOX — UPN of the mailbox (e.g. ops-team@yourorg.com)

    Azure AD App permissions needed:
        Mail.Read (Application permission, admin-consented)

    Optional constructor args:
        folder        — mail folder name (default "Inbox")
        max_messages  — max emails to fetch (default 200)
        allowed_roles / sensitivity_level
    """

    def __init__(
        self,
        mailbox: Optional[str] = None,
        folder: str = "Inbox",
        max_messages: int = 200,
        allowed_roles: Optional[list[str]] = None,
        sensitivity_level: str = "internal",
        **kwargs,
    ):
        self.mailbox = mailbox or settings.ms_mailbox
        self.folder = folder
        self.max_messages = max_messages
        self.allowed_roles = allowed_roles or ["operations", "compliance", "admin"]
        self.sensitivity_level = sensitivity_level

    def fetch(self) -> list[RawDocument]:
        if not self.mailbox:
            raise ValueError(
                "OutlookConnector: MS_MAILBOX is not set. "
                "Add it to your .env file."
            )

        token = _get_ms_graph_token()
        headers = {"Authorization": f"Bearer {token}"}
        base_url = f"https://graph.microsoft.com/v1.0/users/{self.mailbox}"

        documents: list[RawDocument] = []
        # $select limits the fields returned for faster API calls
        params = {
            "$select": "id,subject,from,receivedDateTime,body,importance,sensitivity",
            "$top": 50,
            "$orderby": "receivedDateTime desc",
        }
        url = f"{base_url}/mailFolders/{self.folder}/messages"

        logger.info("OutlookConnector: fetching from %s/%s", self.mailbox, self.folder)

        with httpx.Client(timeout=30) as client:
            while url and len(documents) < self.max_messages:
                resp = client.get(url, headers=headers, params=params)
                resp.raise_for_status()
                data = resp.json()
                messages = data.get("value", [])

                for msg in messages:
                    try:
                        msg_id = msg.get("id", str(uuid.uuid4()))
                        subject = msg.get("subject", "(No subject)")
                        from_addr = (msg.get("from") or {}).get("emailAddress", {})
                        sender = f"{from_addr.get('name', '')} <{from_addr.get('address', '')}>"
                        received = msg.get("receivedDateTime", "")
                        body_html = (msg.get("body") or {}).get("content", "")
                        plain_body = _html_to_text(body_html)

                        content = (
                            f"Subject: {subject}\n"
                            f"From: {sender}\n"
                            f"Received: {received}\n"
                            f"Folder: {self.folder}\n\n"
                            f"{plain_body}"
                        )

                        # Map Outlook sensitivity to internal levels
                        outlook_sensitivity = msg.get("sensitivity", "normal")
                        sensitivity_map = {
                            "normal": "internal",
                            "personal": "internal",
                            "private": "confidential",
                            "confidential": "restricted",
                        }
                        s_level = sensitivity_map.get(outlook_sensitivity, self.sensitivity_level)

                        doc = RawDocument(
                            source_id=f"outlook-{msg_id}",
                            source_name=f"Email: {subject[:80]}",
                            source_path=f"outlook://{self.mailbox}/{self.folder}/{msg_id}",
                            source_type="email",
                            content=content,
                            metadata={
                                "message_id": msg_id,
                                "subject": subject,
                                "from": sender,
                                "received": received,
                                "folder": self.folder,
                                "mailbox": self.mailbox,
                            },
                            pages=[{"page_num": 1, "text": content}],
                            allowed_roles=self.allowed_roles,
                            sensitivity_level=s_level,
                        )
                        documents.append(doc)
                    except Exception as exc:
                        logger.warning("Outlook message error: %s", exc)

                # Follow Graph API @odata.nextLink for pagination
                url = data.get("@odata.nextLink")
                params = {}  # nextLink already contains params

        logger.info("OutlookConnector: fetched %d emails", len(documents))
        return documents


# ── SharePoint Connector ──────────────────────────────────────────────────────

class SharePointConnector(BaseConnector):
    """
    Fetches files from a SharePoint document library via Microsoft Graph API.
    Downloads supported file types (pdf, txt, md, docx) and parses them
    using the project's existing file parsing utilities.

    Required .env vars:
        MS_TENANT_ID, MS_CLIENT_ID, MS_CLIENT_SECRET
        MS_SHAREPOINT_SITE_ID  — e.g. yourorg.sharepoint.com,site-guid,web-guid
        MS_SHAREPOINT_DRIVE_ID — drive GUID (Documents library)

    Azure AD App permissions needed:
        Sites.Read.All (Application permission, admin-consented)
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def __init__(
        self,
        site_id: Optional[str] = None,
        drive_id: Optional[str] = None,
        folder_path: str = "root",
        allowed_roles: Optional[list[str]] = None,
        sensitivity_level: str = "internal",
        **kwargs,
    ):
        self.site_id = site_id or settings.ms_sharepoint_site_id
        self.drive_id = drive_id or settings.ms_sharepoint_drive_id
        self.folder_path = folder_path
        self.allowed_roles = allowed_roles or ["analyst", "operations", "compliance", "admin"]
        self.sensitivity_level = sensitivity_level

    def _check_credentials(self):
        missing = [
            v for v, val in [
                ("MS_SHAREPOINT_SITE_ID", self.site_id),
                ("MS_SHAREPOINT_DRIVE_ID", self.drive_id),
            ] if not val
        ]
        if missing:
            raise ValueError(
                f"SharePointConnector: missing env vars: {', '.join(missing)}. "
                "Add them to your .env file."
            )

    def _parse_file_bytes(self, name: str, content_bytes: bytes) -> str:
        """Pass raw bytes through the project's parsers to extract plain text."""
        ext = "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""

        if ext == ".txt" or ext == ".md":
            return content_bytes.decode("utf-8", errors="replace")

        if ext == ".pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=content_bytes, filetype="pdf")
                return "\n".join(page.get_text() for page in doc)
            except Exception as exc:
                logger.warning("PDF parse error for %s: %s", name, exc)
                return ""

        if ext == ".docx":
            try:
                import docx
                doc = docx.Document(io.BytesIO(content_bytes))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception as exc:
                logger.warning("DOCX parse error for %s: %s", name, exc)
                return ""

        return ""

    def fetch(self) -> list[RawDocument]:
        self._check_credentials()
        token = _get_ms_graph_token()
        headers = {"Authorization": f"Bearer {token}"}

        base = (
            f"https://graph.microsoft.com/v1.0/sites/{self.site_id}"
            f"/drives/{self.drive_id}"
        )
        folder_url = (
            f"{base}/root/children"
            if self.folder_path == "root"
            else f"{base}/root:/{self.folder_path}:/children"
        )

        documents: list[RawDocument] = []
        logger.info(
            "SharePointConnector: fetching from site=%s drive=%s folder=%s",
            self.site_id, self.drive_id, self.folder_path,
        )

        with httpx.Client(timeout=60) as client:
            # Paginate through drive items
            list_url: Optional[str] = folder_url
            while list_url:
                resp = client.get(list_url, headers=headers)
                resp.raise_for_status()
                data = resp.json()
                items = data.get("value", [])

                for item in items:
                    name = item.get("name", "")
                    ext = "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""
                    if ext not in self.SUPPORTED_EXTENSIONS:
                        continue

                    download_url = item.get("@microsoft.graph.downloadUrl")
                    if not download_url:
                        continue

                    item_id = item.get("id", str(uuid.uuid4()))
                    web_url = item.get("webUrl", "")
                    size = item.get("size", 0)

                    try:
                        file_resp = client.get(download_url, timeout=120)
                        file_resp.raise_for_status()
                        plain_text = self._parse_file_bytes(name, file_resp.content)

                        if not plain_text.strip():
                            continue

                        doc = RawDocument(
                            source_id=f"sharepoint-{item_id}",
                            source_name=f"SharePoint: {name}",
                            source_path=web_url or f"sharepoint://{self.site_id}/{name}",
                            source_type="sharepoint",
                            content=plain_text,
                            metadata={
                                "item_id": item_id,
                                "file_name": name,
                                "size_bytes": size,
                                "web_url": web_url,
                                "site_id": self.site_id,
                                "drive_id": self.drive_id,
                            },
                            pages=[{"page_num": 1, "text": plain_text}],
                            allowed_roles=self.allowed_roles,
                            sensitivity_level=self.sensitivity_level,
                        )
                        documents.append(doc)
                        logger.info("SharePoint: ingested %s (%d bytes)", name, size)
                    except Exception as exc:
                        logger.warning("SharePoint file error for %s: %s", name, exc)

                list_url = data.get("@odata.nextLink")

        logger.info("SharePointConnector: fetched %d documents", len(documents))
        return documents
