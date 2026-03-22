"""File system connector: PDF, TXT, DOCX, Markdown."""
import hashlib
import io
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pdfplumber
from PIL import Image

from src.ingestion.connectors.base import BaseConnector, RawDocument

# Lazy-import pytesseract to avoid hard failure if not installed
_tesseract_available: Optional[bool] = None

def _ocr_available() -> bool:
    global _tesseract_available
    if _tesseract_available is None:
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            _tesseract_available = True
        except Exception:
            _tesseract_available = False
    return _tesseract_available


def _ocr_image(pil_image: Image.Image) -> str:
    """Run Tesseract OCR on a PIL image and return extracted text."""
    import pytesseract
    return pytesseract.image_to_string(pil_image, lang="eng").strip()


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

# Pages with fewer chars than this are treated as image/scanned pages
SCANNED_PAGE_THRESHOLD = 50


class FileConnector(BaseConnector):
    """
    Connects to a local directory or a single file and ingests documents.
    Supports PDF (with table + OCR extraction), plain text, markdown, and docx.
    """

    def __init__(
        self,
        path: str,
        allowed_roles: Optional[list[str]] = None,
        sensitivity_level: str = "internal",
        recursive: bool = True,
    ):
        self.path = Path(path)
        self.allowed_roles = allowed_roles or ["analyst", "compliance", "operations", "admin"]
        self.sensitivity_level = sensitivity_level
        self.recursive = recursive

    def fetch(self) -> list[RawDocument]:
        docs = []
        if self.path.is_file():
            files = [self.path]
        elif self.path.is_dir():
            pattern = "**/*" if self.recursive else "*"
            files = [
                f for f in self.path.glob(pattern)
                if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
            ]
        else:
            raise FileNotFoundError(f"Path not found: {self.path}")

        for file_path in files:
            try:
                doc = self._process_file(file_path)
                if doc:
                    docs.append(doc)
            except Exception as e:
                print(f"[FileConnector] Error processing {file_path}: {e}")

        return docs

    def _process_file(self, file_path: Path) -> Optional[RawDocument]:
        ext = file_path.suffix.lower()
        source_id = hashlib.sha256(str(file_path.resolve()).encode()).hexdigest()[:16]

        if ext == ".pdf":
            return self._process_pdf(file_path, source_id)
        elif ext in {".txt", ".md"}:
            return self._process_text(file_path, source_id)
        elif ext == ".docx":
            return self._process_docx(file_path, source_id)
        return None

    def _process_pdf(self, file_path: Path, source_id: str) -> RawDocument:
        pages = []
        full_text_parts = []
        ocr_used = False

        with fitz.open(str(file_path)) as pdf:
            for page_num, page in enumerate(pdf, start=1):
                # --- Native text extraction ---
                text = page.get_text("text").strip()

                # --- OCR fallback for scanned / image-only pages ---
                if len(text) < SCANNED_PAGE_THRESHOLD and _ocr_available():
                    ocr_text = self._ocr_page(page)
                    if ocr_text:
                        text = f"[OCR] {ocr_text}"
                        ocr_used = True

                # --- Extract and OCR embedded images on text pages ---
                image_texts = []
                if _ocr_available():
                    image_texts = self._ocr_embedded_images(page, pdf)

                # Combine text and image descriptions for this page
                page_content = text
                if image_texts:
                    joined = "\n".join(f"[Image {i+1}] {t}" for i, t in enumerate(image_texts))
                    page_content = f"{text}\n\n{joined}" if text else joined

                if page_content:
                    pages.append({"page_num": page_num, "text": page_content})
                    full_text_parts.append(f"[Page {page_num}]\n{page_content}")

        # --- Table extraction via pdfplumber ---
        tables = []
        try:
            with pdfplumber.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    for t_idx, table in enumerate(page.extract_tables() or []):
                        if not table:
                            continue
                        md = _table_to_markdown(table)
                        tables.append({
                            "table_index": t_idx,
                            "page_num": page_num,
                            "markdown": md,
                        })
        except Exception:
            pass

        metadata: dict = {"file_size": file_path.stat().st_size}
        if ocr_used:
            metadata["ocr"] = True

        return RawDocument(
            source_id=source_id,
            source_name=file_path.name,
            source_path=str(file_path.resolve()),
            source_type="pdf",
            content="\n\n".join(full_text_parts),
            metadata=metadata,
            pages=pages,
            tables=tables,
            allowed_roles=self.allowed_roles,
            sensitivity_level=self.sensitivity_level,
        )

    def _ocr_page(self, page: fitz.Page) -> str:
        """Render a full page to an image and OCR it."""
        try:
            mat = fitz.Matrix(2.0, 2.0)  # 2× zoom for better OCR accuracy
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            return _ocr_image(img)
        except Exception:
            return ""

    def _ocr_embedded_images(self, page: fitz.Page, pdf: fitz.Document) -> list[str]:
        """Extract embedded images from a page and OCR each one."""
        results = []
        try:
            for img_info in page.get_images(full=True):
                xref = img_info[0]
                base_image = pdf.extract_image(xref)
                img_bytes = base_image.get("image")
                if not img_bytes:
                    continue
                img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                # Skip tiny images (icons, decorations)
                if img.width < 100 or img.height < 100:
                    continue
                text = _ocr_image(img)
                if text and len(text) > 20:  # Only keep meaningful OCR output
                    results.append(text)
        except Exception:
            pass
        return results

    def _process_text(self, file_path: Path, source_id: str) -> RawDocument:
        content = file_path.read_text(encoding="utf-8", errors="replace")
        return RawDocument(
            source_id=source_id,
            source_name=file_path.name,
            source_path=str(file_path.resolve()),
            source_type="txt" if file_path.suffix == ".txt" else "markdown",
            content=content,
            metadata={"file_size": file_path.stat().st_size},
            allowed_roles=self.allowed_roles,
            sensitivity_level=self.sensitivity_level,
        )

    def _process_docx(self, file_path: Path, source_id: str) -> RawDocument:
        try:
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
        except Exception as e:
            content = f"[DOCX parse error: {e}]"

        return RawDocument(
            source_id=source_id,
            source_name=file_path.name,
            source_path=str(file_path.resolve()),
            source_type="docx",
            content=content,
            metadata={"file_size": file_path.stat().st_size},
            allowed_roles=self.allowed_roles,
            sensitivity_level=self.sensitivity_level,
        )


def _table_to_markdown(table: list[list]) -> str:
    """Convert a pdfplumber table (list of rows) to Markdown."""
    if not table:
        return ""
    rows = []
    for i, row in enumerate(table):
        cleaned = [str(cell).strip() if cell is not None else "" for cell in row]
        rows.append("| " + " | ".join(cleaned) + " |")
        if i == 0:
            rows.append("|" + "|".join(["---"] * len(cleaned)) + "|")
    return "\n".join(rows)
